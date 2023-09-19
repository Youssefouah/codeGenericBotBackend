from fastapi import APIRouter, UploadFile, File, Body
import requests
from docx import Document
from docx2pdf import convert
import os
from llama_index import download_loader

from utils import generate_filename

router = APIRouter()


@router.post("/image")
async def upload_image(image: list[UploadFile] = File(...)):
    total_content = await image[0].read()
    total_content = total_content.decode("utf-8")
    return {"content": str(total_content)}


@router.post("/url")
async def process_url(url: str = Body(..., embed=True)):
    UnstructuredURLLoader = download_loader("UnstructuredURLLoader")
    print(url)
    urls = [url]

    loader = UnstructuredURLLoader(
        urls=urls, continue_on_failure=False, headers={"User-Agent": "value"}
    )
    content = loader.load()
    return {"content": content[0].text}


@router.post("/export/{doc_type}")
async def export_pdf(doc_type: str, advice: dict = Body(..., embed=True)):
    file_path = generate_filename(f"1.{doc_type}")
    if doc_type == "pdf":
        document = Document()
        for key, value in advice.items():
            document.add_heading(key)
            document.add_paragraph(value)
        document.save(f"./static/{file_path.replace('.pdf', '.docx')}")
        convert(
            f"./static/{file_path.replace('.pdf', '.docx')}", f"./static/{file_path}"
        )
        os.remove(f"./static/{file_path.replace('.pdf', '.docx')}")
    if doc_type == "txt":
        text_content = ""
        for key, value in advice.items():
            text_content += f"{key}\n{value}\n\n"
        with open(f"./static/{file_path}", "w") as file:
            file.write(text_content)
    if doc_type == "docx":
        document = Document()
        for key, value in advice.items():
            document.add_heading(key)
            document.add_paragraph(value)
        document.save(f"./static/{file_path}")
    return {"path": file_path}
